#!/usr/bin/env python3
"""
Generate styled Word documents from markdown files.
AWS-branded styling: Amazon Ember 11pt, AWS Orange (#FF9900) accents,
status badges, and clean table formatting.

Usage:
    python3 generate_styled_docx.py input.md [-o output.docx] [-l en|ko] [--footer "custom footer text"]
"""

import argparse
import atexit
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

# === Colors ===
AWS_ORANGE = "FF9900"
AWS_ORANGE_LIGHT = "FFF5E6"
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BADGE_HIGH = "DC3545"
BADGE_MEDIUM = "FF9900"
BADGE_LOW = "28A745"
BADGE_ON_ROADMAP = "0073BB"
BADGE_NOT_STARTED = "6C757D"
TABLE_HEADER_BG = "232F3E"  # AWS dark navy
TABLE_ALT_ROW = "F8F9FA"
CODE_BG = "F0F0F0"

# Meridian (classic Amazon narrative) palette - pure black & white
MERIDIAN_TEXT = RGBColor(0x00, 0x00, 0x00)
MERIDIAN_TABLE_HEADER_BG = "D9D9D9"  # light gray header fill
MERIDIAN_BORDER = "808080"           # gray cell borders
MERIDIAN_RULE = "000000"             # black heading underline


# === Theme (reconfigured per --style) ========================================
# The inline renderers below are module-level functions and can't see the
# builder instance, so the body font / size / color and inline-code color live
# here and are reset by configure_theme() at the start of each conversion.
BODY_FONT = "Amazon Ember"
BODY_SIZE = 11
BODY_COLOR = DARK_GRAY
INLINE_CODE_COLOR = RGBColor(0xC7, 0x25, 0x4E)


def configure_theme(style):
    """Set the module-level font/size/color theme used by the inline renderers."""
    global BODY_FONT, BODY_SIZE, BODY_COLOR, INLINE_CODE_COLOR
    if style == "meridian":
        BODY_FONT = "Calibri"
        BODY_SIZE = 10.5
        BODY_COLOR = MERIDIAN_TEXT
        INLINE_CODE_COLOR = RGBColor(0x33, 0x33, 0x33)
    else:
        BODY_FONT = "Amazon Ember"
        BODY_SIZE = 11
        BODY_COLOR = DARK_GRAY
        INLINE_CODE_COLOR = RGBColor(0xC7, 0x25, 0x4E)


# -- Helpers ------------------------------------------------------------------

def set_cell_shading(cell, color):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = parse_xml(f"<w:tcBorders {nsdecls('w')}/>")
        tcPr.append(borders)
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val[0]}" '
                f'w:space="0" w:color="{val[1]}"/>'
            )
            existing = borders.find(qn(f"w:{side}"))
            if existing is not None:
                borders.remove(existing)
            borders.append(el)


def add_paragraph_border(paragraph, side="bottom", sz="6", color=AWS_ORANGE, space="4"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:{side} w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_run_shading(run, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    run._element.get_or_add_rPr().append(shading)


def set_paragraph_spacing(paragraph, before=None, after=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def set_table_full_width(tbl):
    """Set table to fill page width with proper cell margins.

    Forces fixed layout so a long cell value can't blow the table past the
    page margin. Also disables row no-wrap and tells cells to break long
    tokens so URLs and code stay inside the column.
    """
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl._tbl.insert(0, tblPr)
    # Replace existing width/layout/margin specs
    for tag in ('w:tblW', 'w:tblCellMar', 'w:tblLayout'):
        existing = tblPr.find(qn(tag))
        if existing is not None:
            tblPr.remove(existing)
    tblPr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="60" w:type="dxa"/>'
        f'<w:left w:w="120" w:type="dxa"/>'
        f'<w:bottom w:w="60" w:type="dxa"/>'
        f'<w:right w:w="120" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))
    # Per-row: disable cantSplit/noWrap; per-cell: ensure tcW set + allow wrap
    for row in tbl.rows:
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is not None:
            for child in list(trPr):
                if child.tag in (qn('w:cantSplit'), qn('w:noWrap')):
                    trPr.remove(child)
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            # Remove any noWrap inside the cell — long tokens must wrap
            for nw in tcPr.findall(qn('w:noWrap')):
                tcPr.remove(nw)


# -- Inline markdown rendering ------------------------------------------------

INLINE_RE = re.compile(
    r'(\*\*(.+?)\*\*)'       # group 1,2 = bold
    r'|(\*(.+?)\*)'          # group 3,4 = italic
    r'|(`(.+?)`)'            # group 5,6 = code
    r'|(\[(.+?)\]\((.+?)\))' # group 7,8,9 = link
    r'|(\^(\d+))'            # group 10,11 = footnote ref (^N)
)


def add_formatted_text(paragraph, text, base_size=None, base_bold=False,
                       base_color=None, base_italic=False):
    """Render inline markdown (bold, italic, code, link, ^N footnote) into a paragraph."""
    if base_size is None:
        base_size = BODY_SIZE
    if base_color is None:
        base_color = BODY_COLOR
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _plain_run(paragraph, text[pos:m.start()], base_size, base_bold,
                       base_color, base_italic)

        if m.group(1):  # **bold**
            r = paragraph.add_run(m.group(2))
            r.font.name = BODY_FONT
            r.font.size = Pt(base_size)
            r.font.bold = True
            r.font.italic = base_italic
            r.font.color.rgb = base_color
        elif m.group(3):  # *italic*
            r = paragraph.add_run(m.group(4))
            r.font.name = BODY_FONT
            r.font.size = Pt(base_size)
            r.font.italic = True
            r.font.color.rgb = base_color
        elif m.group(5):  # `code`
            r = paragraph.add_run(m.group(6))
            r.font.name = "Consolas"
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = INLINE_CODE_COLOR
            add_run_shading(r, CODE_BG)
        elif m.group(7):  # [text](url)
            link_text = m.group(8)
            link_url = m.group(9)
            add_hyperlink(paragraph, link_text, link_url, font_size=base_size)
        elif m.group(10):  # ^N footnote reference
            r = paragraph.add_run(m.group(11))
            r.font.name = BODY_FONT
            r.font.size = Pt(max(base_size - 2, 7))
            r.font.color.rgb = base_color
            r.font.superscript = True

        pos = m.end()

    if pos < len(text):
        _plain_run(paragraph, text[pos:], base_size, base_bold, base_color, base_italic)


def _plain_run(paragraph, text, size, bold, color, italic=False):
    if size is None:
        size = BODY_SIZE
    r = paragraph.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color


def add_hyperlink(paragraph, text, url, font_name=None, font_size=11):
    """Add a clickable hyperlink to a paragraph."""
    if font_name is None:
        font_name = BODY_FONT
    from xml.sax.saxutils import escape as xml_escape

    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
    )

    sz_val = int(font_size * 2)
    escaped_text = xml_escape(text)
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{sz_val}"/>'
        f'<w:color w:val="0073BB"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{escaped_text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_badge(paragraph, label, bg_color):
    """Insert a small coloured badge (pill) inline."""
    r = paragraph.add_run(f"  {label}  ")
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = WHITE
    add_run_shading(r, bg_color)
    paragraph.add_run(" ")


# -- Badge rules --------------------------------------------------------------

BADGE_RULES_EN = {
    "on the roadmap": ("ON ROADMAP", BADGE_ON_ROADMAP),
    "on our roadmap": ("ON ROADMAP", BADGE_ON_ROADMAP),
    "not today": ("NOT TODAY", BADGE_NOT_STARTED),
    "not directly": ("LIMITED", BADGE_NOT_STARTED),
    "cannot": ("LIMITATION", BADGE_NOT_STARTED),
}

BADGE_RULES_KO = {
    "로드맵에 있": ("ON ROADMAP", BADGE_ON_ROADMAP),
    "로드맵": ("ON ROADMAP", BADGE_ON_ROADMAP),
    "현재는 미지원": ("NOT TODAY", BADGE_NOT_STARTED),
    "현재는 불가": ("NOT TODAY", BADGE_NOT_STARTED),
    "직접적으로는 불가": ("LIMITED", BADGE_NOT_STARTED),
    "접근 불가": ("LIMITATION", BADGE_NOT_STARTED),
}

PRIORITY_BADGE = {
    "High": BADGE_HIGH, "높음": BADGE_HIGH,
    "Medium": BADGE_MEDIUM, "중간": BADGE_MEDIUM,
    "Low": BADGE_LOW, "낮음": BADGE_LOW,
}


# -- Mermaid rendering (optional; falls back to a code box on any failure) ----

# Font stack applied inside rendered SVG/PNG so Korean glyphs don't break.
_MERMAID_FONT_STACK = (
    '"Apple SD Gothic Neo", "Noto Sans KR", "Malgun Gothic", '
    '"Nanum Gothic", -apple-system, "Segoe UI", sans-serif'
)
_MERMAID_CSS = (
    f"svg {{ font-family: {_MERMAID_FONT_STACK}; }}\n"
    ".nodeLabel, .edgeLabel, .titleText, .taskText, .sectionTitle, "
    f"text, span, div, p {{ font-family: {_MERMAID_FONT_STACK} !important; }}\n"
)

# mermaid-cli is only probed once per process.
_MERMAID_CLI_CHECKED = False
_MERMAID_CLI_AVAILABLE = False


def mermaid_cli_available():
    """True if npx is on PATH so mermaid-cli can be invoked. Probed once."""
    global _MERMAID_CLI_CHECKED, _MERMAID_CLI_AVAILABLE
    if not _MERMAID_CLI_CHECKED:
        _MERMAID_CLI_CHECKED = True
        _MERMAID_CLI_AVAILABLE = shutil.which("npx") is not None
    return _MERMAID_CLI_AVAILABLE


def render_mermaid_to_png(source, workdir, index, scale=3):
    """Render one mermaid diagram to a PNG. Returns the file path, or None on failure.

    Rendering to a high-scale PNG keeps the diagram crisp when Word rescales it.
    Any failure (no npx, mmdc error, timeout) returns None so the caller can
    fall back to rendering the mermaid source as a code box.
    """
    if not mermaid_cli_available():
        return None
    try:
        css = os.path.join(workdir, "mermaid.css")
        if not os.path.exists(css):
            with open(css, "w", encoding="utf-8") as fh:
                fh.write(_MERMAID_CSS)
        mmd = os.path.join(workdir, f"fig{index}.mmd")
        png = os.path.join(workdir, f"fig{index}.png")
        with open(mmd, "w", encoding="utf-8") as fh:
            fh.write(source + "\n")
        proc = subprocess.run(
            ["npx", "-y", "-p", "@mermaid-js/mermaid-cli", "mmdc",
             "-i", mmd, "-o", png,
             "--cssFile", css, "-b", "white", "-s", str(scale)],
            capture_output=True, text=True, timeout=180, check=False,
        )
        if os.path.exists(png) and os.path.getsize(png) > 0:
            return png
        detail = proc.stderr[-500:] if proc.stderr else ""
        sys.stderr.write(
            f"  (mermaid {index} render failed; falling back to code box)\n{detail}\n")
    except Exception as e:  # noqa: BLE001 - never let a diagram break the build
        sys.stderr.write(f"  (mermaid {index} render error: {e}; falling back)\n")
    return None


# -- Document builder ---------------------------------------------------------

DEFAULT_MARGINS_CM = {"top": 2.54, "bottom": 2.54, "left": 2.0, "right": 2.0}


class StyledDocxBuilder:
    def __init__(self, lang="en", footer_text=None, margins=None, style="aws"):
        self.doc = Document()
        self.lang = lang
        self.style = style
        self.footer_text = footer_text
        self.margins = {**DEFAULT_MARGINS_CM, **(margins or {})}
        self.badge_rules = BADGE_RULES_EN if lang == "en" else BADGE_RULES_KO
        self.accent = MERIDIAN_TEXT if style == "meridian" else RGBColor(0xFF, 0x99, 0x00)
        self.footnotes = {}
        self.extra_notes = []
        # Mermaid diagram rendering (lazy temp dir, per-doc counter)
        self._mermaid_workdir = None
        self._mermaid_count = 0
        configure_theme(style)
        self._setup_styles()

    def _get_mermaid_workdir(self):
        if self._mermaid_workdir is None:
            self._mermaid_workdir = tempfile.mkdtemp(prefix="docx-mermaid-")
            atexit.register(shutil.rmtree, self._mermaid_workdir, ignore_errors=True)
        return self._mermaid_workdir

    def _setup_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(BODY_SIZE)
        normal.font.color.rgb = BODY_COLOR
        pf = normal.paragraph_format
        if self.style == "meridian":
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            pf.space_after = Pt(6)
            pf.line_spacing = 1.15

        for section in self.doc.sections:
            section.top_margin = Cm(self.margins["top"])
            section.bottom_margin = Cm(self.margins["bottom"])
            section.left_margin = Cm(self.margins["left"])
            section.right_margin = Cm(self.margins["right"])

        if self.style == "meridian":
            # Section headers are body-size bold black - no large heading sizes
            sizes = {1: BODY_SIZE, 2: BODY_SIZE, 3: BODY_SIZE}
            befores = {1: 8, 2: 8, 3: 6}
            afters = {1: 2, 2: 2, 3: 2}
            head_color = MERIDIAN_TEXT
        else:
            sizes = {1: 22, 2: 15, 3: 12}
            befores = {1: 12, 2: 18, 3: 12}  # H1 gets 12pt so the title doesn't hug the top margin
            afters = {1: 6, 2: 6, 3: 4}
            head_color = DARK_GRAY
        for lvl in (1, 2, 3):
            hs = self.doc.styles[f"Heading {lvl}"]
            hs.font.name = BODY_FONT
            hs.font.size = Pt(sizes[lvl])
            hs.font.bold = True
            hs.font.color.rgb = head_color
            hs.paragraph_format.space_before = Pt(befores[lvl])
            hs.paragraph_format.space_after = Pt(afters[lvl])

    # -- public API -----------------------------------------------------------

    def build(self, md_path, out_path):
        with open(md_path, encoding="utf-8") as f:
            raw = f.read()
        raw = self._preprocess_html(raw)
        lines = raw.splitlines()
        self.footnotes, self.extra_notes, skip_lines = self._extract_footnotes(lines)
        i = 0
        in_code_block = False
        code_block_lines = []
        code_block_lang = ""
        while i < len(lines):
            if i in skip_lines:
                i += 1
                continue
            line = lines[i]

            # -- Fenced code block handling --
            fence_m = re.match(r"^(`{3,})(.*)?$", line.strip())
            if fence_m:
                if not in_code_block:
                    in_code_block = True
                    code_block_lang = (fence_m.group(2) or "").strip()
                    code_block_lines = []
                    i += 1
                    continue
                else:
                    # End of code block — render it.
                    # A mermaid block becomes an embedded image when mermaid-cli
                    # is available; otherwise it falls back to the code box.
                    if code_block_lang.lower() == "mermaid" and \
                            self._render_mermaid_block(code_block_lines):
                        pass
                    else:
                        self._render_code_block(code_block_lines, code_block_lang)
                    in_code_block = False
                    code_block_lines = []
                    code_block_lang = ""
                    i += 1
                    continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # render horizontal rules as thin separator
            if re.match(r"^---+$", line.strip()):
                self._render_horizontal_rule()
                i += 1
                continue

            # -- Heading --
            hm = re.match(r"^(#{1,6})\s+(.+)$", line)
            if hm:
                level = len(hm.group(1))
                text = hm.group(2)
                # H4-H6: render as bold paragraph (python-docx only supports H1-H3 natively)
                if level <= 3:
                    p = self.doc.add_heading(level=level)
                    p.clear()
                    if self.style == "meridian":
                        add_formatted_text(p, text, base_size=BODY_SIZE,
                                           base_bold=True, base_color=BODY_COLOR)
                        if level <= 2:
                            add_paragraph_border(p, side="bottom", sz="6",
                                                 color=MERIDIAN_RULE, space="1")
                    else:
                        size = {1: 22, 2: 15, 3: 12}[level]
                        add_formatted_text(p, text, base_size=size,
                                           base_bold=True, base_color=DARK_GRAY)
                        if level == 1:
                            add_paragraph_border(p, sz="8", color=AWS_ORANGE, space="6")
                        elif level == 2:
                            add_paragraph_border(p)
                else:
                    p = self.doc.add_paragraph()
                    if self.style == "meridian":
                        add_formatted_text(p, text, base_size=BODY_SIZE,
                                           base_bold=True, base_color=BODY_COLOR)
                        set_paragraph_spacing(p, before=8, after=2)
                    else:
                        size = {4: 11, 5: 11, 6: 10}.get(level, 11)
                        add_formatted_text(p, text, base_size=size,
                                           base_bold=True, base_color=DARK_GRAY)
                        set_paragraph_spacing(p, before=10, after=4)
                i += 1
                continue

            # -- Blockquote block --
            if line.startswith(">"):
                bq_lines = []
                while i < len(lines) and lines[i].startswith(">"):
                    bq_lines.append(re.sub(r"^>\s?", "", lines[i]))
                    i += 1
                joined = " ".join(bq_lines)
                if "Key Takeaways" in joined or "핵심 요약" in joined:
                    self._render_key_takeaways(bq_lines)
                elif self._is_footer_metadata(bq_lines):
                    self._render_footer_metadata(bq_lines)
                else:
                    self._render_blockquote(bq_lines)
                continue

            # -- Table --
            if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
                table_lines = []
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._render_table(table_lines)
                continue

            # -- Numbered list --
            nm = re.match(r"^(\d+)\.\s+(.+)$", line)
            if nm:
                items = []
                while i < len(lines):
                    nm2 = re.match(r"^(\d+)\.\s+(.+)$", lines[i])
                    if nm2 or lines[i].startswith("   "):
                        items.append(lines[i])
                        i += 1
                    else:
                        break
                self._render_numbered_list(items)
                continue

            # -- Bullet list --
            bm = re.match(r"^(\s*)- (.+)$", line)
            if bm:
                items = []
                while i < len(lines):
                    bm2 = re.match(r"^(\s*)- (.+)$", lines[i])
                    if bm2:
                        items.append((len(bm2.group(1)) // 2, bm2.group(2)))
                        i += 1
                    else:
                        break
                self._render_bullet_list(items)
                continue

            # -- Metadata lines (Author, Date, Attendees) --
            meta_m = re.match(r"^\*\*(.+?):\*\*\s*(.+)$", line)
            if meta_m:
                p = self.doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=2)
                r = p.add_run(f"{meta_m.group(1)}: ")
                r.font.name = BODY_FONT
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                add_formatted_text(p, meta_m.group(2), base_size=10,
                                   base_color=RGBColor(0x66, 0x66, 0x66))
                i += 1
                continue

            # -- Normal paragraph --
            if line.strip():
                p = self.doc.add_paragraph()
                add_formatted_text(p, line)
                self._maybe_add_badge(p, line)
                i += 1
                continue

            # blank line
            i += 1

        # -- Endnotes + Footer --
        self._render_endnotes()
        self._add_footer()
        self.doc.save(out_path)
        print(f"  -> {out_path}")

    def _preprocess_html(self, text):
        """Convert common GitHub-README HTML to markdown equivalents and strip the rest.

        - <a href="URL">TEXT</a>           → [TEXT](URL)
        - <strong>X</strong>, <b>X</b>     → **X**
        - <h1>...</h1> .. <h6>...</h6>     → # X .. ###### X (align attr ignored)
        - <img ...>                        → removed (badges, logos)
        - <details>, </details>            → removed (wrapper only)
        - <summary>X</summary>             → **X** (acts as a heading-ish label)
        - <p>, </p>, <sub>, </sub>, <br>   → removed/whitespace
        Other tags are stripped but content is preserved.
        """
        # Strip <img ...> entirely (self-closing or not)
        text = re.sub(r"<img\b[^>]*/?>", "", text, flags=re.IGNORECASE)
        # <a href="URL">TEXT</a> → [TEXT](URL)
        text = re.sub(
            r'<a\b[^>]*\bhref="([^"]+)"[^>]*>(.*?)</a>',
            r"[\2](\1)",
            text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        # <strong>X</strong>, <b>X</b> → **X**
        text = re.sub(r"<(strong|b)\b[^>]*>(.*?)</\1>", r"**\2**", text, flags=re.IGNORECASE | re.DOTALL)
        # <hN>X</hN> → markdown heading
        for n in range(1, 7):
            text = re.sub(
                rf"<h{n}\b[^>]*>(.*?)</h{n}>",
                lambda m, n=n: f"\n{'#' * n} {m.group(1).strip()}\n",
                text,
                flags=re.IGNORECASE | re.DOTALL,
            )
        # <summary>X</summary> → **X** on its own line (collapsible section heading)
        text = re.sub(
            r"<summary\b[^>]*>(.*?)</summary>",
            lambda m: f"\n**{m.group(1).strip()}**\n",
            text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        # <br> / <br/> → newline
        text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
        # Drop wrapper-only tags entirely (open and close)
        for tag in ("details", "p", "sub", "sup", "div", "span", "center"):
            text = re.sub(rf"</?{tag}\b[^>]*>", "", text, flags=re.IGNORECASE)
        return text

    def _extract_footnotes(self, lines):
        """Scan for '**Notes:**' / 'Notes:' sections and extract footnote defs.

        Returns ({n: text}, [extra unnumbered notes], set of line indices to skip).
        Only matches exact 'Notes:' headers — 'Important Notes:' etc. are left alone.
        Consumes the whole Notes block (contiguous bullets + blank lines) until a
        non-bullet, non-blank line — so bullets without a ^N prefix still get lifted
        into the endnotes block instead of being rendered as body.
        """
        footnotes = {}
        extras = []
        skip = set()
        header_re = re.compile(r'^\s*(?:\*\*Notes:\*\*|Notes:)\s*$')
        num_item_re = re.compile(
            r'^\s*-\s+\*\*\^(\d+)(?:\s+([^:*]+?))?:?\*\*\s*(.*)$'
        )
        bullet_re = re.compile(r'^\s*-\s+(.+)$')
        i = 0
        while i < len(lines):
            if header_re.match(lines[i]):
                block_indices = [i]
                j = i + 1
                found_any = False
                while j < len(lines):
                    mm = num_item_re.match(lines[j])
                    if mm:
                        n = int(mm.group(1))
                        label = (mm.group(2) or "").strip()
                        body = mm.group(3).strip()
                        footnotes[n] = f"{label}: {body}" if label else body
                        block_indices.append(j)
                        found_any = True
                        j += 1
                        continue
                    bm = bullet_re.match(lines[j])
                    if bm:
                        extras.append(bm.group(1).strip())
                        block_indices.append(j)
                        found_any = True
                        j += 1
                        continue
                    if lines[j].strip() == "":
                        block_indices.append(j)
                        j += 1
                        continue
                    break
                if found_any:
                    skip.update(block_indices)
                    i = j
                    continue
            i += 1
        return footnotes, extras, skip

    def _render_endnotes(self):
        """Render extracted footnotes as a styled endnotes block at document end."""
        if not self.footnotes and not self.extra_notes:
            return

        gray = RGBColor(0x66, 0x66, 0x66)

        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=18, after=4)
        add_paragraph_border(p, side="top", sz="4", color="D0D0D0", space="6")

        for n in sorted(self.footnotes):
            p = self.doc.add_paragraph()
            set_paragraph_spacing(p, before=1, after=1)
            r = p.add_run(str(n))
            r.font.name = BODY_FONT
            r.font.size = Pt(8)
            r.font.superscript = True
            r.font.color.rgb = gray
            p.add_run(" ").font.size = Pt(9)
            add_formatted_text(p, self.footnotes[n], base_size=9,
                               base_color=gray, base_italic=True)

        for note in self.extra_notes:
            p = self.doc.add_paragraph()
            set_paragraph_spacing(p, before=1, after=1)
            r = p.add_run("  •  ")
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.color.rgb = gray
            add_formatted_text(p, note, base_size=9,
                               base_color=gray, base_italic=True)

    # -- Renderers ------------------------------------------------------------

    def _render_key_takeaways_plain(self, bq_lines):
        """Meridian: Key Takeaways as a plain bold heading + bullets (no box/orange)."""
        title = "Key Takeaways" if self.lang == "en" else "핵심 요약"
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=8, after=2)
        add_formatted_text(p, title, base_size=BODY_SIZE, base_bold=True,
                           base_color=BODY_COLOR)
        add_paragraph_border(p, side="bottom", sz="6", color=MERIDIAN_RULE, space="1")
        for line in bq_lines:
            line = line.strip()
            if not line or "Key Takeaways" in line or "핵심 요약" in line:
                continue
            line = line.removeprefix("- ")
            bp = self.doc.add_paragraph()
            bp.paragraph_format.left_indent = Cm(0.6)
            set_paragraph_spacing(bp, before=1, after=1)
            r = bp.add_run("•  ")
            r.font.name = BODY_FONT
            r.font.size = Pt(BODY_SIZE)
            r.font.color.rgb = BODY_COLOR
            add_formatted_text(bp, line, base_size=BODY_SIZE)

    def _render_key_takeaways(self, bq_lines):
        """Render Key Takeaways as an AWS-orange accented box using a 1-cell table."""
        if self.style == "meridian":
            self._render_key_takeaways_plain(bq_lines)
            return
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = 1  # center
        set_table_full_width(tbl)
        cell = tbl.cell(0, 0)

        set_cell_borders(cell,
                         left=("24", AWS_ORANGE),
                         top=("4", "E0E0E0"),
                         bottom=("4", "E0E0E0"),
                         right=("4", "E0E0E0"))
        set_cell_shading(cell, AWS_ORANGE_LIGHT)
        cell.paragraphs[0].clear()

        first = True
        for line in bq_lines:
            line = line.strip()
            if not line:
                continue

            # Title line
            if "Key Takeaways" in line or "핵심 요약" in line:
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                title = "Key Takeaways" if self.lang == "en" else "핵심 요약"
                r = p.add_run(title)
                r.font.name = BODY_FONT
                r.font.size = Pt(13)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)
                set_paragraph_spacing(p, before=4, after=6)
                continue

            # Bullet point
            if line.startswith("- "):
                p = cell.add_paragraph()
                first = False
                r = p.add_run("  \u2022  ")
                r.font.name = BODY_FONT
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)
                r.font.bold = True
                add_formatted_text(p, line[2:], base_size=10)
                set_paragraph_spacing(p, before=2, after=2)

        # Spacing after the table
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    def _render_blockquote(self, bq_lines):
        """Render a general blockquote with a gray left border."""
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = 1
        set_table_full_width(tbl)
        cell = tbl.cell(0, 0)

        set_cell_borders(cell,
                         left=("18", "CCCCCC"),
                         top=("2", "E8E8E8"),
                         bottom=("2", "E8E8E8"),
                         right=("2", "E8E8E8"))
        set_cell_shading(cell, "F9F9F9")
        cell.paragraphs[0].clear()

        first = True
        for line in bq_lines:
            line = line.strip()
            if not line:
                continue
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            add_formatted_text(p, line, base_size=10, base_color=RGBColor(0x55, 0x55, 0x55))
            set_paragraph_spacing(p, before=2, after=2)

        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    def _render_mermaid_block(self, code_lines):
        """Render a mermaid block as a centered, embedded image.

        Returns True if the image was embedded, False if the caller should fall
        back to rendering the mermaid source as a code box.
        """
        source = "\n".join(code_lines).strip()
        if not source:
            return False
        self._mermaid_count += 1
        png = render_mermaid_to_png(
            source, self._get_mermaid_workdir(), self._mermaid_count)
        if not png:
            return False
        try:
            from PIL import Image  # optional; used only to size the image
            with Image.open(png) as im:
                px_w, px_h = im.size
        except Exception:  # noqa: BLE001 - Pillow optional
            px_w, px_h = (0, 0)

        section = self.doc.sections[0]
        usable_emu = section.page_width - section.left_margin - section.right_margin
        usable_cm = usable_emu / 360000.0
        # Keep diagrams to ~85% of text width so they read as figures, and cap
        # height so a tall flowchart doesn't overflow a page.
        target_cm = usable_cm * 0.85
        width_cm = target_cm
        if px_w and px_h:
            # 96 dpi baseline; -s scale in the renderer only sharpens, not resizes
            native_cm = px_w / 96.0 * 2.54 / 3.0  # scale=3 was used
            width_cm = min(target_cm, native_cm)
            max_h_cm = 20.0  # leave room within an A4/Letter text column
            if width_cm * (px_h / px_w) > max_h_cm:
                width_cm = max_h_cm * (px_w / px_h)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=6, after=2)
        run = p.add_run()
        try:
            run.add_picture(png, width=Cm(width_cm))
        except Exception as e:  # noqa: BLE001 - never break the build on embed
            sys.stderr.write(f"  (mermaid embed failed: {e}; falling back)\n")
            # roll back the empty paragraph we just added
            p._element.getparent().remove(p._element)
            return False
        # trailing spacer
        sp = self.doc.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=4)
        return True

    def _render_code_block(self, code_lines, lang=""):
        """Render a fenced code block as a shaded monospace box."""
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = 1
        set_table_full_width(tbl)
        cell = tbl.cell(0, 0)

        set_cell_borders(cell,
                         top=("2", "D0D0D0"),
                         bottom=("2", "D0D0D0"),
                         left=("2", "D0D0D0"),
                         right=("2", "D0D0D0"))
        set_cell_shading(cell, "F5F5F5")
        cell.paragraphs[0].clear()

        # Optional language label
        first = True
        if lang:
            p = cell.paragraphs[0]
            first = False
            r = p.add_run(lang.upper())
            r.font.name = "Consolas"
            r.font.size = Pt(7)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            set_paragraph_spacing(p, before=2, after=0)

        # Pick a font size that keeps the longest line inside the page width.
        # Consolas at 9pt ≈ 1.95 mm per char; usable width = pageWidth - L/R margins.
        section = self.doc.sections[0]
        usable_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
        max_cols = max((len(line) for line in code_lines), default=0)
        # Empirical char widths in cm for Consolas at given pt
        char_cm = {9: 0.195, 8: 0.173, 7: 0.152, 6: 0.130}
        code_pt = 9
        # Account for cell padding (~0.4cm each side) and a small safety margin
        budget = max(usable_cm - 0.9, 6.0)
        for pt in (9, 8, 7, 6):
            if max_cols * char_cm[pt] <= budget:
                code_pt = pt
                break
        else:
            code_pt = 6  # very long lines — fall back to smallest

        for line in code_lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            r = p.add_run(line if line else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(code_pt)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            set_paragraph_spacing(p, before=0, after=0)

        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=4)

    def _render_table(self, table_lines):
        """Render a markdown table with styled header and alternating rows."""
        rows_raw = []
        for tl in table_lines:
            if re.match(r"^\|[\s\-:|]+\|$", tl.strip()):
                continue
            cells = [c.strip() for c in tl.strip().strip("|").split("|")]
            rows_raw.append(cells)

        if not rows_raw:
            return

        n_cols = len(rows_raw[0])
        n_rows = len(rows_raw)

        tbl = self.doc.add_table(rows=n_rows, cols=n_cols)
        tbl.alignment = 1
        tbl.autofit = True
        set_table_full_width(tbl)

        # Style header row
        header_text_color = MERIDIAN_TEXT if self.style == "meridian" else WHITE
        header_bg = MERIDIAN_TABLE_HEADER_BG if self.style == "meridian" else TABLE_HEADER_BG
        for j, val in enumerate(rows_raw[0]):
            cell = tbl.cell(0, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_text(p, val, base_size=10, base_bold=True, base_color=header_text_color)
            set_cell_shading(cell, header_bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_paragraph_spacing(p, before=4, after=4)

        # Data rows
        for i in range(1, n_rows):
            for j in range(n_cols):
                cell = tbl.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]

                val = rows_raw[i][j].strip() if j < len(rows_raw[i]) else ""

                if val in PRIORITY_BADGE and self.style != "meridian":
                    add_badge(p, val, PRIORITY_BADGE[val])
                else:
                    add_formatted_text(p, val, base_size=10)

                if i % 2 == 0 and self.style != "meridian":
                    set_cell_shading(cell, TABLE_ALT_ROW)

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_paragraph_spacing(p, before=3, after=3)

        # Set thin borders on all cells
        border_color = MERIDIAN_BORDER if self.style == "meridian" else "D0D0D0"
        border_sz = "4" if self.style == "meridian" else "2"
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_borders(cell,
                                 top=(border_sz, border_color),
                                 bottom=(border_sz, border_color),
                                 left=(border_sz, border_color),
                                 right=(border_sz, border_color))

        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=4)

    def _render_numbered_list(self, items):
        for item in items:
            nm = re.match(r"^(\d+)\.\s+(.+)$", item)
            if nm:
                p = self.doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=2)
                r = p.add_run(f"{nm.group(1)}.  ")
                r.font.name = BODY_FONT
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = self.accent
                add_formatted_text(p, nm.group(2))
                self._maybe_add_badge(p, nm.group(2))
            elif item.strip().startswith("- "):
                p = self.doc.add_paragraph()
                set_paragraph_spacing(p, before=1, after=1)
                p.paragraph_format.left_indent = Cm(1.2)
                r = p.add_run("\u2022  ")
                r.font.name = BODY_FONT
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                add_formatted_text(p, item.strip()[2:], base_size=10)
            elif item.strip():
                # Continuation text under a numbered item
                p = self.doc.add_paragraph()
                set_paragraph_spacing(p, before=1, after=1)
                p.paragraph_format.left_indent = Cm(0.8)
                add_formatted_text(p, item.strip(), base_size=10)

    def _render_bullet_list(self, items):
        for indent_level, text in items:
            p = self.doc.add_paragraph()
            left = Cm(0.6 * indent_level)
            p.paragraph_format.left_indent = left
            set_paragraph_spacing(p, before=2, after=2)

            bullet_char = "\u2022" if indent_level == 0 else "\u25E6"
            r = p.add_run(f"  {bullet_char}  ")
            r.font.name = BODY_FONT
            r.font.size = Pt(11)
            if indent_level == 0:
                r.font.color.rgb = self.accent
            else:
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            add_formatted_text(p, text)
            self._maybe_add_badge(p, text)

    def _render_horizontal_rule(self):
        """Render a horizontal rule as a thin gray separator line."""
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=6, after=6)
        add_paragraph_border(p, side="bottom", sz="4", color="D0D0D0", space="0")

    def _is_footer_metadata(self, bq_lines):
        """Check if blockquote is document footer metadata."""
        joined = " ".join(bq_lines)
        markers = ["Author:", "작성자:", "Source:", "출처:",
                    "Last updated:", "최종 업데이트:"]
        return sum(1 for m in markers if m in joined) >= 2

    def _render_footer_metadata(self, bq_lines):
        """Render footer metadata with subtle, compact styling."""
        # Thin top separator
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=16, after=4)
        add_paragraph_border(p, side="bottom", sz="4", color="D0D0D0", space="4")

        for line in bq_lines:
            line = line.strip()
            if not line:
                continue
            p = self.doc.add_paragraph()
            set_paragraph_spacing(p, before=1, after=1)
            # Split on first colon to highlight the label
            colon_idx = line.find(":")
            if colon_idx > 0:
                label = line[:colon_idx + 1]
                value = line[colon_idx + 1:].strip()
                r = p.add_run(label + " ")
                r.font.name = BODY_FONT
                r.font.size = Pt(8)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                add_formatted_text(p, value, base_size=8,
                                   base_color=RGBColor(0x99, 0x99, 0x99))
            else:
                r = p.add_run(line)
                r.font.name = BODY_FONT
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def _maybe_add_badge(self, paragraph, text):
        """Check badge rules and append a badge if text matches."""
        if self.style == "meridian":
            return
        for trigger, (label, color) in self.badge_rules.items():
            if trigger.lower() in text.lower():
                paragraph.add_run("  ")
                add_badge(paragraph, label, color)
                break

    def _add_footer(self):
        """Add a footer credit line (styled per document style)."""
        section = self.doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]

        if self.style == "meridian":
            self._add_meridian_footer(section, fp)
            return

        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=24, after=4)
        add_paragraph_border(p, side="top", sz="4", color="D0D0D0", space="8")

        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.footer_text:
            txt = self.footer_text
        else:
            today = date.today().isoformat()  # noqa: DTZ011 - local calendar date is intended
            txt = f"{today}  |  Confidential"
        r = fp.add_run(txt)
        r.font.name = BODY_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def _add_meridian_footer(self, section, fp):
        """Meridian footer: 'Amazon Confidential' left, 'Page X of Y' right."""
        usable = section.page_width - section.left_margin - section.right_margin
        fp.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)

        left_txt = self.footer_text if self.footer_text else "Amazon Confidential"

        def _style(run):
            run.font.name = BODY_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = MERIDIAN_TEXT

        _style(fp.add_run(left_txt))
        _style(fp.add_run("\t"))
        _style(fp.add_run("Page "))
        self._add_field(fp, "PAGE", _style)
        _style(fp.add_run(" of "))
        self._add_field(fp, "NUMPAGES", _style)

    def _add_field(self, paragraph, instr, style_fn):
        """Append a Word field (e.g. PAGE, NUMPAGES) as a styled run."""
        run = paragraph.add_run()
        style_fn(run)
        r = run._r
        begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instr_el = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> {instr} </w:instrText>'
        )
        end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        r.append(begin)
        r.append(instr_el)
        r.append(end)


# -- Main ---------------------------------------------------------------------

def detect_lang(md_path):
    """Auto-detect language from filename convention (-ko suffix)."""
    base = os.path.splitext(os.path.basename(md_path))[0]
    if base.endswith("-ko"):
        return "ko"
    return "en"


def main():
    parser = argparse.ArgumentParser(
        description="Convert markdown files to styled Word documents with AWS branding."
    )
    parser.add_argument("input", nargs="+", help="Input markdown file path(s)")
    parser.add_argument("-o", "--output", help="Output .docx file path (only valid with single input file)")
    parser.add_argument("-l", "--lang", choices=["en", "ko"], default=None,
                        help="Language for badge rules and labels (default: auto-detect from filename)")
    parser.add_argument("--footer", default=None,
                        help="Custom footer text (default: auto-generated with date)")
    parser.add_argument("-s", "--style", choices=["aws", "meridian"], default="aws",
                        help="Document style: 'aws' (branded, default) or 'meridian' "
                             "(classic Amazon narrative - Calibri, black & white)")
    parser.add_argument("--margin-top", type=float, default=None,
                        help="Top margin in cm (default: per style)")
    parser.add_argument("--margin-bottom", type=float, default=None,
                        help="Bottom margin in cm (default: per style)")
    parser.add_argument("--margin-left", type=float, default=None,
                        help="Left margin in cm (default: per style)")
    parser.add_argument("--margin-right", type=float, default=None,
                        help="Right margin in cm (default: per style)")

    args = parser.parse_args()

    if args.output and len(args.input) > 1:
        print("Error: -o/--output can only be used with a single input file.", file=sys.stderr)
        sys.exit(1)

    if args.style == "meridian":
        style_defaults = {"top": 1.27, "bottom": 1.27, "left": 1.27, "right": 1.27}
    else:
        style_defaults = dict(DEFAULT_MARGINS_CM)
    margins = {
        "top": args.margin_top if args.margin_top is not None else style_defaults["top"],
        "bottom": args.margin_bottom if args.margin_bottom is not None else style_defaults["bottom"],
        "left": args.margin_left if args.margin_left is not None else style_defaults["left"],
        "right": args.margin_right if args.margin_right is not None else style_defaults["right"],
    }

    for md_path in args.input:
        if not os.path.isfile(md_path):
            print(f"Error: Input file not found: {md_path}", file=sys.stderr)
            sys.exit(1)

        out_path = args.output if args.output else os.path.splitext(md_path)[0] + ".docx"
        lang = args.lang if args.lang else detect_lang(md_path)

        print(f"Generating styled Word document (lang={lang}, style={args.style})...")
        builder = StyledDocxBuilder(lang=lang, footer_text=args.footer,
                                    margins=margins, style=args.style)
        builder.build(md_path, out_path)

    print("Done!")


if __name__ == "__main__":
    main()
