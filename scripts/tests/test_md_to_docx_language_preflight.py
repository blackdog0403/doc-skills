"""Regression tests for md-to-docx generated-label language preflight."""

import re
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
CONVERTER = ROOT / "scripts" / "generate_styled_docx.py"
HANGUL = re.compile(r"[가-힣ㄱ-ㅎㅏ-ㅣ]")
QUESTION = "Which language should the generated document labels use"


def document_text(path):
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(parts)


class GeneratedLabelLanguageTests(unittest.TestCase):
    def convert(self, directory, language):
        markdown = directory / "english-report-ko.md"
        output = directory / f"labels-{language}.docx"
        markdown.write_text(
            "# English Source Report\n\n"
            "> **Key Takeaways**\n"
            "> - Source content remains in English.\n\n"
            "## Details\n\nThe Markdown body must not be translated.\n",
            encoding="utf-8",
        )
        result = subprocess.run(
            [
                sys.executable,
                str(CONVERTER),
                str(markdown),
                "-o",
                str(output),
                "-l",
                language,
                "--author",
                "Example Author",
                "--toc",
                "--page-numbers",
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
        return document_text(output)

    def test_explicit_english_generates_only_english_labels(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            text = self.convert(Path(temp_dir), "en")
        self.assertIn("Table of Contents", text)
        self.assertIn("Author: Example Author", text)
        self.assertIn("Key Takeaways", text)
        self.assertIn("Page 1 of 1", text)
        self.assertIn("The Markdown body must not be translated.", text)
        self.assertIsNone(HANGUL.search(text), text)

    def test_explicit_korean_generates_korean_labels_without_translating_source(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            text = self.convert(Path(temp_dir), "ko")
        self.assertIn("목차", text)
        self.assertIn("작성자: Example Author", text)
        self.assertIn("핵심 요약", text)
        self.assertIn("페이지 1 / 1", text)
        self.assertIn("The Markdown body must not be translated.", text)


class LanguagePreflightInstructionTests(unittest.TestCase):
    @property
    def paths(self):
        return [
            ROOT / "skills" / "md-to-docx" / "SKILL.md",
            ROOT / "quick" / "md-to-docx" / "SKILL.md",
        ]

    def test_both_variants_require_explicit_generated_label_language(self):
        for path in self.paths:
            with self.subTest(path=path):
                instructions = path.read_text(encoding="utf-8")
                self.assertIn(QUESTION, instructions)
                self.assertIn("does not translate or rewrite the Markdown", instructions)
                self.assertIn("default", instructions)
                self.assertIn("Never infer", instructions)
                self.assertIn("filename suffix", instructions)

    def test_slash_variant_always_passes_normalized_language(self):
        instructions = (ROOT / "skills" / "md-to-docx" / "SKILL.md").read_text(
            encoding="utf-8"
        )
        self.assertIn('LANGUAGE="<en-or-ko-from-preflight>"', instructions)
        self.assertIn('-l "$LANGUAGE"', instructions)
        self.assertIn("Always invoke the helper with exactly one explicit", instructions)

    def test_quick_language_input_is_required_without_silent_default(self):
        instructions = (ROOT / "quick" / "md-to-docx" / "SKILL.md").read_text(
            encoding="utf-8"
        )
        # Slice to the next input, not specifically `footer`: inputs added between
        # `language` and `footer` must not drag their own `default:` into this check.
        language_input = instructions.split("  - name: language", 1)[1].split(
            "\n  - name:", 1
        )[0]
        self.assertIn("required: true", language_input)
        self.assertNotIn("default:", language_input)

    def test_quick_normalizes_and_passes_language_explicitly(self):
        instructions = (ROOT / "quick" / "md-to-docx" / "SKILL.md").read_text(
            encoding="utf-8"
        )
        self.assertIn('"default": "en"', instructions)
        self.assertIn('"english": "en"', instructions)
        self.assertIn('"korean": "ko"', instructions)
    def test_quick_python_snippet_is_syntactically_valid(self):
        instructions = (ROOT / "quick" / "md-to-docx" / "SKILL.md").read_text(
            encoding="utf-8"
        )
        snippet = instructions.split("```python", 1)[1].split("```", 1)[0]
        compile(snippet, "quick/md-to-docx/SKILL.md", "exec")

        self.assertIn('cmd.extend(["-l", language])', instructions)
        self.assertNotIn('cmd.extend(["-l", language or "en"])', instructions)


if __name__ == "__main__":
    unittest.main()
