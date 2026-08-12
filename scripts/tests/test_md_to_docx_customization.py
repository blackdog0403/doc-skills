"""Regression tests for Markdown-to-DOCX document customization."""

import base64
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_ORIENT

ROOT = Path(__file__).resolve().parents[2]
CONVERTER = ROOT / "scripts" / "generate_styled_docx.py"
ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def package_text(path, member):
    """Return all XML text nodes from one DOCX package member."""
    with zipfile.ZipFile(path) as package:
        root = ElementTree.fromstring(package.read(member))
    return "".join(root.itertext())


class DocxCustomizationTests(unittest.TestCase):
    def convert(self, directory, *options):
        markdown = directory / "input.md"
        output = directory / "output.docx"
        markdown.write_text(
            "```markdown\n# Example_Title\n```\n\n"
            "# Source Report\n\n## Overview\n\nDocument body.\n",
            encoding="utf-8",
        )
        command = [
            sys.executable,
            str(CONVERTER),
            str(markdown),
            "-o",
            str(output),
            *options,
        ]
        result = subprocess.run(command, capture_output=True, text=True, check=False)
        self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
        self.assertTrue(output.exists())
        return output

    def test_all_customizations_are_written_to_docx(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            directory = Path(temp_dir)
            logo = directory / "logo.png"
            logo.write_bytes(ONE_PIXEL_PNG)
            output = self.convert(
                directory,
                "--footer",
                "Architecture | Internal",
                "--title-page",
                "--title",
                "Architecture Review",
                "--subtitle",
                "Executive summary",
                "--author",
                "Example Author",
                "--team",
                "Solutions Architecture",
                "--version",
                "2.1",
                "--classification",
                "Internal",
                "--toc",
                "--page-numbers",
                "--header",
                "Review Draft",
                "--logo",
                str(logo),
                "--page-size",
                "a4",
                "--orientation",
                "landscape",
                "--margin-top",
                "1.5",
                "--margin-bottom",
                "1.6",
                "--margin-left",
                "1.7",
                "--margin-right",
                "1.8",
            )

            document = Document(output)
            paragraph_text = "\n".join(p.text for p in document.paragraphs)
            self.assertIn("Architecture Review", paragraph_text)
            self.assertIn("Executive summary", paragraph_text)
            self.assertIn("Author: Example Author", paragraph_text)
            self.assertIn("Team: Solutions Architecture", paragraph_text)
            self.assertIn("Version: 2.1", paragraph_text)
            self.assertIn("Classification: Internal", paragraph_text)
            self.assertIn("Table of Contents", paragraph_text)

            properties = document.core_properties
            self.assertEqual(properties.title, "Architecture Review")
            self.assertEqual(properties.subject, "Executive summary")
            self.assertEqual(properties.author, "Example Author")
            self.assertEqual(properties.category, "Solutions Architecture")
            self.assertIn("version:2.1", properties.keywords)
            self.assertEqual(properties.comments, "Internal")

            section = document.sections[0]
            self.assertEqual(section.orientation, WD_ORIENT.LANDSCAPE)
            self.assertAlmostEqual(section.page_width.cm, 29.7, places=1)
            self.assertAlmostEqual(section.page_height.cm, 21.0, places=1)
            self.assertAlmostEqual(section.top_margin.cm, 1.5, places=1)
            self.assertAlmostEqual(section.bottom_margin.cm, 1.6, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 1.7, places=1)
            self.assertAlmostEqual(section.right_margin.cm, 1.8, places=1)
            self.assertTrue(section.different_first_page_header_footer)
            self.assertIn("Review Draft", section.header.paragraphs[0].text)

            with zipfile.ZipFile(output) as package:
                names = package.namelist()
                document_xml = package_text(output, "word/document.xml")
                footer_xml = "".join(
                    package_text(output, name)
                    for name in names
                    if name.startswith("word/footer") and name.endswith(".xml")
                )
                settings_xml = package.read("word/settings.xml").decode("utf-8")
            self.assertTrue(any(name.startswith("word/media/") for name in names))
            self.assertIn(' TOC \\o "1-3" \\h \\z \\u ', document_xml)
            self.assertIn(" PAGE ", footer_xml)
            self.assertIn(" NUMPAGES ", footer_xml)
            self.assertIn("Architecture | Internal", footer_xml)
            self.assertIn("updateFields", settings_xml)

    def test_metadata_enables_title_page_and_infers_title(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = self.convert(Path(temp_dir), "--author", "Example Author")
            document = Document(output)
        self.assertEqual(document.core_properties.title, "Source Report")
        self.assertEqual(document.core_properties.author, "Example Author")
        self.assertGreaterEqual(
            sum(paragraph.text == "Source Report" for paragraph in document.paragraphs),
            2,
        )
        self.assertTrue(document.sections[0].different_first_page_header_footer)

    def test_missing_logo_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            directory = Path(temp_dir)
            markdown = directory / "input.md"
            markdown.write_text("# Report\n", encoding="utf-8")
            result = subprocess.run(
                [
                    sys.executable,
                    str(CONVERTER),
                    str(markdown),
                    "--logo",
                    str(directory / "missing.png"),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("Logo file not found", result.stderr)


class CustomizationInstructionTests(unittest.TestCase):
    def test_both_skill_variants_prompt_for_all_customizations(self):
        paths = [
            ROOT / "skills" / "md-to-docx" / "SKILL.md",
            ROOT / "quick" / "md-to-docx" / "SKILL.md",
        ]
        expected = [
            "Any other document customization?",
            "title page",
            "table of contents",
            "Page X of Y",
            "header text/logo",
            "Letter/A4/Legal",
            "orientation",
            "margins",
        ]
        for path in paths:
            with self.subTest(path=path):
                instructions = path.read_text(encoding="utf-8")
                for text in expected:
                    self.assertIn(text, instructions)

    def test_slash_skill_maps_every_new_cli_flag(self):
        instructions = (ROOT / "skills" / "md-to-docx" / "SKILL.md").read_text(
            encoding="utf-8"
        )
        for flag in [
            "--title-page",
            "--title",
            "--subtitle",
            "--author",
            "--team",
            "--version",
            "--classification",
            "--toc",
            "--page-numbers",
            "--header",
            "--logo",
            "--page-size",
            "--orientation",
        ]:
            with self.subTest(flag=flag):
                self.assertIn(flag, instructions)

    def test_quick_skill_passes_every_new_cli_flag(self):
        instructions = (ROOT / "quick" / "md-to-docx" / "SKILL.md").read_text(
            encoding="utf-8"
        )
        for flag in [
            "--title-page",
            "--title",
            "--subtitle",
            "--author",
            "--team",
            "--version",
            "--classification",
            "--toc",
            "--page-numbers",
            "--header",
            "--logo",
            "--page-size",
            "--orientation",
        ]:
            with self.subTest(flag=flag):
                self.assertIn(flag, instructions)


if __name__ == "__main__":
    unittest.main()
